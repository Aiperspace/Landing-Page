import base64
import asyncio
import hashlib
import io
import json
import logging
import os
import re
import time
from datetime import date
from html import escape
from typing import Any, AsyncGenerator, Callable, Dict, List, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from openai import AzureOpenAI, OpenAI
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from llm_schema import validate_llm_payload
from pdf_compat import extract_text_from_pdf

logger = logging.getLogger("aiper")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class DocumentRequest(BaseModel):
    description: str
    template_type: str
    notes: str = ""
    files: List[str] = []
    section_outline: Optional[List[str]] = Field(
        default=None,
        description="When set, used as section headings instead of the built-in outline for template_type.",
    )


class Section(BaseModel):
    title: str
    content: str


class GeneratedDocument(BaseModel):
    title: str
    sections: List[Section]


TEMPLATE_OUTLINES = {
    "test_procedure": [
        "Purpose",
        "Scope & Applicability",
        "Prerequisites",
        "Safety Considerations",
        "Procedure Steps",
        "Success Criteria",
        "Data Recording",
    ],
    "test_log": [
        "Test Identification",
        "Test Configuration",
        "Execution Log",
        "Deviations & Non-Conformances",
        "Sign-Off & Approval",
    ],
    "subsystem_verification": [
        "Subsystem Identification",
        "Verification Cross-Reference Matrix",
        "Test Campaign Summary",
        "Open Items & Waivers",
        "Compliance Statement",
        "Annexes & Supporting Evidence",
    ],
    "test_report": [
        "Introduction",
        "Applicable and Reference Documents",
        "Definitions and Abbreviations",
        "Test Results",
        "Anomalies",
        "Conclusions",
    ],
    "icd": [
        "Scope",
        "Applicable Documents",
        "Reference Documents",
        "Definitions, Acronyms, Abbreviations",
        "System Overview",
        "Interface Overview",
        "Mechanical / Electrical / Data interfaces",
        "Environmental & verification",
    ],
}


@app.get("/")
def read_root():
    return {"message": "AI Document backend is running"}


MAX_UPLOAD_CONTEXT_CHARS = 30_000
EXTRACTION_CACHE_TTL_SEC = int(os.getenv("EXTRACTION_CACHE_TTL_SEC", "3600"))
PLAN_CACHE_TTL_SEC = int(os.getenv("PLAN_CACHE_TTL_SEC", "1800"))

# Simple in-process caches. Safe default for single service instance; can be replaced with Redis later.
EXTRACTION_CACHE: dict[str, tuple[float, str]] = {}
PLAN_CACHE: dict[str, tuple[float, dict[str, Any]]] = {}

TEMPLATE_PROMPT_SPECS = {
    "test_procedure": {
        "description": "Step-by-step execution document for running a controlled engineering test campaign.",
        "writing_goals": [
            "Use imperative language for executable test steps.",
            "State prerequisites, required tools, and entry conditions before execution.",
            "Define pass/fail criteria with measurable thresholds where available.",
        ],
        "evidence_style": "Prioritize concrete setup details, ordered steps, and explicit verification outcomes.",
        "ecss_checklist": {
            "required_sections": [
                "Purpose",
                "Scope & Applicability",
                "Prerequisites",
                "Safety Considerations",
                "Procedure Steps",
                "Success Criteria",
                "Data Recording",
            ],
            "mandatory_statements": [
                "Safety constraints and hazard controls are stated.",
                "Each major step has a verifiable expected outcome.",
                "Pass/fail criteria are testable and not subjective.",
            ],
        },
    },
    "test_log": {
        "description": "Chronological record of how a test was configured and executed, including deviations.",
        "writing_goals": [
            "Capture what happened in execution order and keep observations traceable.",
            "Record objective outcomes and anomalies without interpretation drift.",
            "Include sign-off context and unresolved issues.",
        ],
        "evidence_style": "Prefer time-ordered facts, measured values, and explicit anomaly records.",
        "ecss_checklist": {
            "required_sections": [
                "Test Identification",
                "Test Configuration",
                "Execution Log",
                "Deviations & Non-Conformances",
                "Sign-Off & Approval",
            ],
            "mandatory_statements": [
                "Execution log reflects the actual performed sequence.",
                "All deviations are captured with impact notes when known.",
                "Approval/sign-off status is explicitly stated.",
            ],
        },
    },
    "subsystem_verification": {
        "description": "Compliance summary mapping subsystem requirements to verification evidence and open items.",
        "writing_goals": [
            "Map requirements to available evidence and verification method.",
            "Separate verified items from open items and waivers.",
            "End with a clear compliance statement bounded by available data.",
        ],
        "evidence_style": "Use requirement IDs, matrix-style traceability, and concise compliance reasoning.",
        "ecss_checklist": {
            "required_sections": [
                "Subsystem Identification",
                "Verification Cross-Reference Matrix",
                "Test Campaign Summary",
                "Open Items & Waivers",
                "Compliance Statement",
                "Annexes & Supporting Evidence",
            ],
            "mandatory_statements": [
                "Verification status is traceable to evidence.",
                "Open items and waivers are explicitly separated from compliant items.",
                "Compliance statement includes scope limitations when data is incomplete.",
            ],
        },
    },
    "test_report": {
        "description": "Formal ECSS-aligned report documenting objectives, results, anomalies, and conclusions.",
        "writing_goals": [
            "Maintain objective, formal aerospace reporting tone.",
            "Summarize results against objectives and acceptance criteria.",
            "Document anomalies, impact, and recommended follow-up actions.",
        ],
        "evidence_style": "Use measured results, observed behavior, and clear references to supplied evidence.",
        "ecss_checklist": {
            "required_sections": [
                "Introduction",
                "Applicable and Reference Documents",
                "Definitions and Abbreviations",
                "Test Results",
                "Anomalies",
                "Conclusions",
            ],
            "mandatory_statements": [
                "Objectives and test context are clearly stated.",
                "Results are linked to acceptance intent when evidence exists.",
                "Conclusions distinguish confirmed findings from missing data.",
            ],
        },
    },
    "icd": {
        "description": "Interface Control Document defining mechanical, electrical, and data interfaces between systems.",
        "writing_goals": [
            "Specify interface boundaries and assumptions unambiguously.",
            "Describe interface characteristics with verifiable constraints.",
            "State verification approach for each major interface domain.",
        ],
        "evidence_style": "Prefer concrete interface parameters, constraints, and compatibility conditions.",
        "ecss_checklist": {
            "required_sections": [
                "Scope",
                "Applicable Documents",
                "Reference Documents",
                "Definitions, Acronyms, Abbreviations",
                "System Overview",
                "Interface Overview",
                "Mechanical / Electrical / Data interfaces",
                "Environmental & verification",
            ],
            "mandatory_statements": [
                "Interface assumptions and boundaries are explicit.",
                "Mechanical/electrical/data constraints are stated where evidence exists.",
                "Verification method or status is provided for key interface aspects.",
            ],
        },
    },
}

LLM_COMPARE_SYSTEM_PROMPT = (
    "You are a product compliance analyst. "
    "Compare two product specification documents and return a JSON object that strictly matches this schema: "
    '{"<Parameter_Name>": {"A": "Value in Product A", "B": "Value in Product B", '
    '"Status": "OK" | "Different" | "Missing"}, "Summary": "Overall summary text"}. '
    "Rules: no arrays, no extra top-level keys, include all relevant parameters you can find "
    "and return multiple parameter entries, A/B may be strings or numbers, "
    "and Status must be exactly OK, Different, or Missing."
)

COMPONENT_TYPES = [
    "obc",
    "star_tracker",
    "gnss_gps",
    "imu_gyroscope",
    "reaction_wheel",
    "magnetorquer",
    "magnetometer",
    "sun_sensor",
    "earth_horizon_sensor",
    "vhf_uhf_transceiver",
    "tmtc",
    "rf_transmitter",
    "antenna",
    "eps",
    "pcdu_pdu_pmu",
    "acu",
    "battery",
    "solar_panel",
    "solar_cell",
    "sada",
    "payload",
    "structure",
    "unknown",
]

UNIVERSAL_PARAMETER_WHITELIST = [
    "input_voltage_min",
    "input_voltage_nominal",
    "input_voltage_max",
    "average_current",
    "peak_current",
    "inrush_current",
    "average_power",
    "peak_power",
    "interface_type",
    "protocol",
    "data_rate",
    "baud_rate",
    "logic_level",
    "connector",
    "mass",
    "dimensions",
    "operating_temperature_min",
    "operating_temperature_max",
    "survival_temperature_min",
    "survival_temperature_max",
    "radiation_tid",
    "vibration_qualification",
    "shock_qualification",
]

PAIR_PARAMETER_WHITELISTS: Dict[str, list[str]] = {
    "obc__star_tracker": [
        "interface_type",
        "protocol",
        "data_rate",
        "baud_rate",
        "logic_level",
        "differential_or_single_ended",
        "connector",
        "pinout_reference",
        "command_set",
        "telemetry_format",
        "output_rate",
        "measurement_latency",
        "timestamp_support",
        "pps_support",
        "sync_support",
        "input_voltage_min",
        "input_voltage_nominal",
        "input_voltage_max",
        "average_power",
        "peak_power",
        "boresight_reference",
    ],
    "obc__gnss_gps": [
        "supported_constellations",
        "gnss_protocol_messages",
        "update_rate",
        "pps_output",
        "time_accuracy",
        "interface_type",
        "protocol",
        "data_rate",
        "connector",
        "input_voltage_nominal",
        "average_power",
    ],
    "obc__imu_gyroscope": [
        "interface_type",
        "protocol",
        "data_rate",
        "sample_rate",
        "output_update_rate",
        "measurement_latency",
        "timestamp_support",
        "sync_support",
        "input_voltage_nominal",
        "average_power",
    ],
    "obc__reaction_wheel": [
        "command_interface",
        "telemetry_interface",
        "control_update_rate",
        "fault_output_lines",
        "inhibit_lines",
        "state_after_reset",
        "input_voltage_nominal",
        "peak_current",
        "peak_power",
        "mass",
        "dimensions",
    ],
    "obc__magnetorquer": [
        "command_interface_if_smart",
        "driver_requirement_if_passive",
        "duty_cycle",
        "input_voltage_nominal",
        "peak_current",
        "thermal_dissipation",
        "axis_orientation",
    ],
    "obc__magnetometer": [
        "interface_type",
        "protocol",
        "data_rate",
        "output_rate",
        "timestamp_support",
        "measurement_range",
        "accuracy",
        "noise",
    ],
    "obc__sun_sensor": [
        "interface_type",
        "protocol",
        "output_rate",
        "latency",
        "field_of_view",
        "accuracy",
        "axis_or_boresight",
    ],
    "obc__earth_horizon_sensor": [
        "interface_type",
        "protocol",
        "output_rate",
        "latency",
        "field_of_view",
        "boresight_axis_reference",
    ],
    "obc__vhf_uhf_transceiver": [
        "obc_interface",
        "protocol",
        "uplink_data_rate",
        "downlink_data_rate",
        "packet_format",
        "command_set",
        "telemetry_format",
        "rx_current",
        "tx_current",
        "peak_tx_current",
    ],
    "obc__tmtc": [
        "obc_side_interface",
        "protocol",
        "packet_format",
        "uplink_data_rate",
        "downlink_data_rate",
        "buffer_size",
        "sync_clock_requirement",
    ],
    "obc__rf_transmitter": [
        "input_data_interface",
        "input_data_rate",
        "packet_format",
        "rf_output_power",
        "frequency_band",
        "modulation",
    ],
    "obc__eps": [
        "main_bus_voltage",
        "regulated_output_voltages",
        "number_of_output_channels",
        "max_current_per_channel",
        "current_limit_per_channel",
        "trip_threshold",
        "trip_delay",
        "latch_off_reset_behavior",
        "power_sequencing",
        "telemetry_interface",
        "command_interface",
    ],
    "obc__pcdu_pdu_pmu": [
        "main_bus_voltage",
        "regulated_output_voltages",
        "max_current_per_channel",
        "trip_threshold",
        "trip_delay",
        "power_sequencing",
        "telemetry_interface",
    ],
    "obc__payload": [
        "command_interface",
        "telemetry_interface",
        "science_data_interface",
        "data_rate",
        "burst_data_rate",
        "daily_data_volume",
        "timestamp_requirement",
        "latency_tolerance",
        "input_voltage_range",
        "average_power",
        "peak_power",
    ],
    "obc__structure": [
        "dimensions",
        "mass",
        "mounting_interface",
        "mounting_hole_pattern",
        "connector_access",
        "cable_routing_volume",
        "payload_fov_constraints",
    ],
    "vhf_uhf_transceiver__antenna": [
        "frequency_range",
        "frequency_band",
        "impedance",
        "rf_output_power",
        "max_rf_power",
        "antenna_connector",
        "connector_type",
        "polarization",
        "gain",
        "vswr",
    ],
    "rf_transmitter__antenna": [
        "frequency_band",
        "impedance",
        "rf_output_power",
        "max_rf_power",
        "rf_connector",
        "connector_type",
        "polarization",
        "vswr",
    ],
    "tmtc__rf_transmitter": [
        "packet_format",
        "input_data_interface",
        "input_data_rate",
        "modulation",
        "coding",
        "sync_clock_requirement",
    ],
    "gnss_gps__antenna": [
        "supported_constellations",
        "antenna_interface",
        "rf_connector",
        "frequency_band",
        "active_or_passive_antenna_support",
        "antenna_supply_voltage",
    ],
    "battery__acu": [
        "chemistry",
        "voltage_min",
        "voltage_nominal",
        "voltage_max",
        "max_charge_current",
        "charge_voltage",
        "charge_algorithm",
        "protection_behavior",
    ],
    "battery__pcdu_pdu_pmu": [
        "voltage_min",
        "voltage_nominal",
        "voltage_max",
        "capacity_ah",
        "max_continuous_discharge_current",
        "peak_discharge_current",
        "battery_interface",
        "protection_behavior",
    ],
    "solar_panel__acu": [
        "voc",
        "vmp",
        "isc",
        "imp",
        "max_input_power",
        "solar_input_voltage_range",
        "solar_input_current",
        "mppt_voltage_range",
        "connector_type",
        "harness_pinout",
    ],
    "solar_cell__solar_panel": [
        "cell_type",
        "efficiency",
        "voc",
        "vmp",
        "isc",
        "imp",
        "temperature_coefficients",
        "radiation_degradation",
    ],
    "sada__solar_panel": [
        "rotation_range",
        "pointing_accuracy",
        "step_resolution",
        "angular_speed",
        "supported_panel_mass",
        "supported_panel_inertia",
        "mechanical_interface",
    ],
    "sada__structure": [
        "mechanical_interface",
        "mounting_hole_pattern",
        "allowable_loads",
        "stiffness",
        "connector_access",
        "cable_routing_volume",
    ],
    "payload__eps": [
        "input_voltage_range",
        "average_power",
        "peak_power",
        "inrush_current",
        "power_by_mode",
        "duty_cycle",
        "trip_threshold",
        "trip_delay",
    ],
    "payload__structure": [
        "dimensions",
        "mass",
        "mounting_pattern",
        "keepout_volume",
        "field_of_view_if_optical",
        "deployment_envelope",
        "connector_location",
    ],
}

NON_COMPARABLE_BY_PAIR: Dict[str, set[str]] = {
    # For OBC<->Star Tracker, mass and outer envelope are informational here;
    # they should not be emitted as direct compatibility verdict rows.
    "obc__star_tracker": {"mass", "dimensions"},
}

INTERFACE_LIKE_PARAMETERS = {
    "interface_type",
    "protocol",
    "connector",
    "obc_interface",
    "command_interface",
    "telemetry_interface",
}


def _pair_key(type_a: str, type_b: str) -> str:
    left, right = sorted([type_a, type_b])
    return f"{left}__{right}"


def _is_empty_value(value: Any) -> bool:
    if value is None:
        return True
    s = str(value).strip()
    if s == "" or s == "—":
        return True

    low = s.lower()
    exact_placeholders = {
        "n/a",
        "na",
        "not provided",
        "not specified",
        "not available",
        "unknown",
        "none",
        "null",
        "tbd",
    }
    if low in exact_placeholders:
        return True

    # Catch placeholders embedded in longer explanatory strings such as:
    # "Not specified (uses RS-422/RS-485 for command/telemetry)"
    placeholder_phrases = (
        "not specified",
        "not provided",
        "not available",
        "unknown",
        "n/a",
        "tbd",
    )
    return any(phrase in low for phrase in placeholder_phrases)


def _tokenize_for_overlap(text: str) -> set[str]:
    low = text.lower()
    # keep technical tokens such as rs-422, can-fd
    chunks = re.split(r"[,/;\n()]+", low)
    tokens: set[str] = set()
    for chunk in chunks:
        for token in chunk.strip().split():
            cleaned = re.sub(r"[^a-z0-9\-]+", "", token)
            if len(cleaned) >= 2:
                tokens.add(cleaned)
    return tokens


def _has_interface_overlap(a: Any, b: Any) -> bool:
    ta = _tokenize_for_overlap(str(a))
    tb = _tokenize_for_overlap(str(b))
    if not ta or not tb:
        return False
    shared = ta.intersection(tb)
    # require at least one meaningful shared token
    return len(shared) > 0


def _apply_compatibility_overrides(result: dict[str, Any], pair_key: str) -> dict[str, Any]:
    # Remove non-comparable parameters for this pair
    non_comparable = NON_COMPARABLE_BY_PAIR.get(pair_key, set())
    for p in list(result.keys()):
        if p in {"Summary", "DetectedComponents"}:
            continue
        if p in non_comparable:
            result.pop(p, None)

    power_conditioning_note_needed = False

    # Promote known interface overlaps from Different -> OK
    for p, entry in list(result.items()):
        if p in {"Summary", "DetectedComponents"} or not isinstance(entry, dict):
            continue
        status = entry.get("Status")
        a_val = entry.get("A")
        b_val = entry.get("B")
        if not isinstance(status, str):
            continue
        missing_a = _is_empty_value(a_val)
        missing_b = _is_empty_value(b_val)

        # Enforce a strict missing-data rule: if one side is absent/placeholder,
        # the row must be Missing (prevents false OK on protocol/connector, etc.).
        if missing_a or missing_b:
            entry["Status"] = "Missing"
            result[p] = entry
            continue

        if p in INTERFACE_LIKE_PARAMETERS and status == "Different":
            if not missing_a and not missing_b and _has_interface_overlap(a_val, b_val):
                entry["Status"] = "OK"
                result[p] = entry

        if p.startswith("input_voltage") and status == "Different":
            power_conditioning_note_needed = True

    if power_conditioning_note_needed:
        summary = str(result.get("Summary", "")).strip()
        note = " Note: input-voltage mismatch may still be integrable with a suitable DC/DC power-conditioning stage."
        if note.strip() not in summary:
            result["Summary"] = f"{summary}{note}".strip()

    return result


def _detect_component_type(client, model_name: str, filename: str, text: str) -> dict[str, Any]:
    prompt = f"""
You classify spacecraft components from technical datasheets.
Choose exactly one component type from this list:
{", ".join(COMPONENT_TYPES)}

Return JSON only:
{{
  "component_type": "<one listed type>",
  "confidence": <number between 0 and 1>,
  "evidence": "short quote or clue"
}}

If uncertain, return "unknown".

Filename: {filename}
Document text:
---
{text[:12000]}
---
"""
    completion = client.chat.completions.create(
        model=model_name,
        temperature=1,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": "You classify aerospace components."},
            {"role": "user", "content": prompt},
        ],
    )
    raw = json.loads(completion.choices[0].message.content or "{}")
    comp_type = str(raw.get("component_type", "unknown")).strip().lower()
    if comp_type not in COMPONENT_TYPES:
        comp_type = "unknown"
    confidence = raw.get("confidence", 0)
    try:
        confidence = max(0.0, min(1.0, float(confidence)))
    except (TypeError, ValueError):
        confidence = 0.0
    evidence = str(raw.get("evidence", "")).strip()[:240]
    return {"component_type": comp_type, "confidence": confidence, "evidence": evidence}


def _get_llm_client():
    provider = os.getenv("LLM_PROVIDER", "openai").lower()
    if provider == "azure":
        azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
        azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
        if not azure_api_key or not azure_endpoint:
            raise HTTPException(
                status_code=500,
                detail=(
                    "Missing Azure settings. Required: AZURE_OPENAI_API_KEY, "
                    "AZURE_OPENAI_ENDPOINT"
                ),
            )
        client = AzureOpenAI(
            api_key=azure_api_key,
            api_version=azure_api_version,
            azure_endpoint=azure_endpoint,
        )
        model_name = os.getenv("OPENAI_MODEL", "gpt-5-mini")
    else:
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            raise HTTPException(
                status_code=500,
                detail="Missing OPENAI_API_KEY. Set it in backend environment variables.",
            )
        client = OpenAI(api_key=openai_api_key)
        model_name = os.getenv("OPENAI_MODEL", "gpt-5-mini")
    return client, model_name


def _extract_docx_text(data: bytes, max_chars: int = 18_000) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts)[:max_chars]


def _extract_xlsx_text(data: bytes, max_rows: int = 120, max_cols: int = 10, max_chars: int = 18_000) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets[:4]:
        lines.append(f"[Sheet: {ws.title}]")
        for i, row in enumerate(ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True)):
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if vals:
                lines.append(" | ".join(vals))
            if i >= max_rows - 1:
                break
        lines.append("")
    return "\n".join(lines)[:max_chars]


def _extract_text_from_upload(filename: str, content_type: str | None, data: bytes) -> str:
    low = filename.lower()
    if low.endswith(".pdf") or (content_type or "").lower() == "application/pdf":
        return _extract_pdf_text(data)
    if low.endswith(".docx"):
        return _extract_docx_text(data)
    if low.endswith(".txt") or low.endswith(".csv"):
        return data.decode("utf-8", errors="ignore")[:18_000]
    if low.endswith(".xlsx"):
        return _extract_xlsx_text(data)
    return ""


def _summarize_image_evidence(client, model_name: str, filename: str, content_type: str | None, data: bytes) -> str:
    mime = content_type or "image/png"
    if not mime.startswith("image/"):
        mime = "image/png"
    b64 = base64.standard_b64encode(data).decode("ascii")
    user_content = [
        {
            "type": "text",
            "text": (
                "This is a test plot/graph/image related to an aerospace test. "
                "Extract only concrete observations visible in the image (values, trends, anomalies, labels, axes). "
                "If nothing readable, say 'Not provided'. Keep it concise."
            ),
        },
        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
    ]
    completion = client.chat.completions.create(
        model=model_name,
        temperature=1,
        messages=[
            {"role": "system", "content": "You extract factual evidence from engineering images."},
            {"role": "user", "content": user_content},
        ],
    )
    text = (completion.choices[0].message.content or "").strip()
    if not text:
        return "Not provided"
    return text[:3_000]


def _resolve_outline(template_type: str, section_outline: Optional[List[str]]) -> list[str]:
    if section_outline:
        outline = [s.strip() for s in section_outline if s.strip()]
        if outline:
            return outline
    return TEMPLATE_OUTLINES.get(template_type, ["Purpose"])


def _cache_get_text(cache: dict[str, tuple[float, str]], key: str, ttl_sec: int) -> Optional[str]:
    row = cache.get(key)
    if not row:
        return None
    ts, value = row
    if time.time() - ts > ttl_sec:
        cache.pop(key, None)
        return None
    return value


def _cache_get_json(cache: dict[str, tuple[float, dict[str, Any]]], key: str, ttl_sec: int) -> Optional[dict[str, Any]]:
    row = cache.get(key)
    if not row:
        return None
    ts, value = row
    if time.time() - ts > ttl_sec:
        cache.pop(key, None)
        return None
    return value


def _pack_context_for_planning(uploaded_context: str, outline: list[str], max_total_chars: int = 16_000) -> str:
    """
    Reduce planning token load by selecting the most relevant evidence blocks per section.
    """
    blocks = [b.strip() for b in uploaded_context.split("\n\n[") if b.strip()]
    normalized_blocks = []
    for i, block in enumerate(blocks):
        b = block if i == 0 else f"[{block}"
        normalized_blocks.append(b)

    if not normalized_blocks:
        return uploaded_context[:max_total_chars]

    selected: list[str] = []
    used = set()
    section_terms = [re.findall(r"[a-z0-9]+", s.lower()) for s in outline]

    for terms in section_terms:
        best_idx = -1
        best_score = -1
        for idx, block in enumerate(normalized_blocks):
            if idx in used:
                continue
            lower = block.lower()
            score = sum(1 for t in terms if t and t in lower)
            if score > best_score:
                best_score = score
                best_idx = idx
        if best_idx >= 0:
            used.add(best_idx)
            selected.append(normalized_blocks[best_idx])

    for idx, block in enumerate(normalized_blocks):
        if idx not in used:
            selected.append(block)
        if len("\n\n".join(selected)) >= max_total_chars:
            break

    return "\n\n".join(selected)[:max_total_chars]


async def _build_uploaded_context(
    files: List[UploadFile],
    client,
    model_name: str,
    on_chunk_ready: Optional[Callable[[str, str], None]] = None,
) -> tuple[list[str], str]:
    async def _process_upload(index: int, upload: UploadFile) -> tuple[int, str, str]:
        name = upload.filename or "upload"
        data = await upload.read()
        if not data:
            return index, name, ""
        digest = hashlib.sha256(data).hexdigest()
        cache_key = f"{name.lower()}|{(upload.content_type or '').lower()}|{digest}|{model_name}"
        cached_text = _cache_get_text(EXTRACTION_CACHE, cache_key, EXTRACTION_CACHE_TTL_SEC)
        if cached_text is not None:
            return index, name, cached_text

        try:
            if _is_probably_image(name, upload.content_type):
                extracted = await asyncio.to_thread(
                    _summarize_image_evidence,
                    client,
                    model_name,
                    name,
                    upload.content_type,
                    data,
                )
                text = f"[{name}] Image evidence summary:\n{extracted}"
            else:
                extracted = await asyncio.to_thread(
                    _extract_text_from_upload,
                    name,
                    upload.content_type,
                    data,
                )
                extracted = extracted.strip()
                if not extracted:
                    text = f"[{name}] No readable content extracted (unsupported type or empty file)."
                else:
                    text = f"[{name}] Extracted content:\n{extracted}"
        except Exception as exc:
            text = f"[{name}] Could not parse file: {exc}"
        EXTRACTION_CACHE[cache_key] = (time.time(), text)
        return index, name, text

    tasks = [asyncio.create_task(_process_upload(i, upload)) for i, upload in enumerate(files)]
    results: list[tuple[int, str, str]] = []
    for done in asyncio.as_completed(tasks):
        result = await done
        results.append(result)
        _, name, text = result
        if text and on_chunk_ready:
            on_chunk_ready(name, text)
    results.sort(key=lambda item: item[0])

    file_names = [name for _, name, _ in results]
    chunks: list[str] = []
    total = 0
    for _, _, text in results:
        if not text:
            continue
        if total + len(text) > MAX_UPLOAD_CONTEXT_CHARS:
            remaining = MAX_UPLOAD_CONTEXT_CHARS - total
            if remaining > 200:
                chunks.append(text[:remaining] + "\n...[truncated]")
            break
        chunks.append(text)
        total += len(text)
    return file_names, "\n\n".join(chunks).strip()


def _parse_markdown_sections(markdown_text: str, outline: list[str], fallback_title: str) -> dict:
    lines = markdown_text.splitlines()
    title = fallback_title
    if lines and lines[0].strip().startswith("# "):
        title = lines[0].strip()[2:].strip() or fallback_title

    sections: list[dict[str, str]] = []
    current_title: Optional[str] = None
    current_lines: list[str] = []
    heading_re = re.compile(r"^##\s+(.+?)\s*$")

    for raw_line in lines:
        line = raw_line.rstrip()
        match = heading_re.match(line.strip())
        if match:
            if current_title:
                sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})
            current_title = match.group(1).strip()
            current_lines = []
        elif current_title is not None:
            current_lines.append(line)

    if current_title:
        sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    if not sections:
        sections = [{"title": s, "content": "Not provided"} for s in outline]

    return {"title": title, "sections": sections}


def _plan_from_context(
    client,
    model_name: str,
    template_type: str,
    outline: list[str],
    template_spec: dict[str, Any],
    today: str,
    description: str,
    notes: str,
    file_names: List[str],
    uploaded_context: str,
    base_plan: Optional[dict[str, Any]] = None,
) -> dict[str, Any]:
    packed_context = _pack_context_for_planning(uploaded_context, outline)
    cache_payload = {
        "template_type": template_type,
        "outline": outline,
        "description": description.strip(),
        "notes": notes.strip(),
        "file_names": file_names,
        "packed_context_hash": hashlib.sha256(packed_context.encode("utf-8", errors="ignore")).hexdigest(),
        "base_plan_hash": hashlib.sha256(
            json.dumps(base_plan or {}, sort_keys=True, ensure_ascii=False).encode("utf-8", errors="ignore")
        ).hexdigest(),
    }
    plan_cache_key = hashlib.sha256(
        json.dumps(cache_payload, sort_keys=True, ensure_ascii=False).encode("utf-8", errors="ignore")
    ).hexdigest()
    cached_plan = _cache_get_json(PLAN_CACHE, plan_cache_key, PLAN_CACHE_TTL_SEC)
    if cached_plan is not None:
        return cached_plan

    checklist_json = json.dumps(template_spec["ecss_checklist"], ensure_ascii=False, indent=2)
    base_plan_text = json.dumps(base_plan, ensure_ascii=False, indent=2) if base_plan else "(none)"
    plan_prompt = f"""
You are a senior aerospace documentation planner.
Build a concise section-by-section plan before writing the final document.

Date: {today}
Template type: {template_type}
Template description: {template_spec["description"]}
Target outline: {outline}
Writing goals: {template_spec["writing_goals"]}
Evidence style: {template_spec["evidence_style"]}
ECSS checklist:
{checklist_json}

Previous draft plan:
{base_plan_text}

User description:
{description}

Additional notes:
{notes or "(none)"}

Uploaded file names:
{file_names or []}

Packed extracted evidence (relevance-filtered):
{packed_context or "(none)"}

Return JSON only with this schema:
{{
  "title": "string",
  "sections": [
    {{
      "title": "string",
      "key_facts": ["string"],
      "evidence_refs": ["string"],
      "data_gaps": ["string"],
      "acceptance_statement": "string"
    }}
  ]
}}
Rules:
- Keep section order exactly equal to target outline.
- key_facts must contain only facts present in the provided inputs/evidence.
- If evidence is missing, add a short item under data_gaps.
- Do not invent requirement IDs, values, dates, or measurements.
- Keep each list concise and non-redundant.
"""
    plan_completion = client.chat.completions.create(
        model=model_name,
        temperature=1,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": "You plan aerospace technical documents with strict evidence grounding."},
            {"role": "user", "content": plan_prompt},
        ],
    )
    plan_raw = plan_completion.choices[0].message.content or "{}"
    try:
        parsed = json.loads(plan_raw)
    except json.JSONDecodeError:
        parsed = {"title": f"{template_type.replace('_', ' ').title()} - {today}", "sections": []}
    PLAN_CACHE[plan_cache_key] = (time.time(), parsed)
    return parsed


def _generate_doc_from_inputs(
    client,
    model_name: str,
    description: str,
    notes: str,
    template_type: str,
    section_outline: Optional[List[str]],
    file_names: List[str],
    uploaded_context: str,
    progress_callback: Optional[Callable[[str], None]] = None,
    stream_draft_callback: Optional[Callable[[str], None]] = None,
    precomputed_early_plan: Optional[dict[str, Any]] = None,
):
    outline = _resolve_outline(template_type, section_outline)
    today = date.today().strftime("%d/%m/%Y")
    template_spec = TEMPLATE_PROMPT_SPECS.get(template_type, TEMPLATE_PROMPT_SPECS["test_procedure"])
    checklist_json = json.dumps(template_spec["ecss_checklist"], ensure_ascii=False, indent=2)
    outline_json = json.dumps(outline, ensure_ascii=False)
    if precomputed_early_plan is not None:
        if progress_callback:
            progress_callback("Refining early plan with full evidence")
        plan_json = _plan_from_context(
            client=client,
            model_name=model_name,
            template_type=template_type,
            outline=outline,
            template_spec=template_spec,
            today=today,
            description=description,
            notes=notes,
            file_names=file_names,
            uploaded_context=uploaded_context,
            base_plan=precomputed_early_plan,
        )
    else:
        if progress_callback:
            progress_callback("Building section plan")
        plan_json = _plan_from_context(
            client=client,
            model_name=model_name,
            template_type=template_type,
            outline=outline,
            template_spec=template_spec,
            today=today,
            description=description,
            notes=notes,
            file_names=file_names,
            uploaded_context=uploaded_context,
        )
    # Hand the writer a prose brief rather than the raw plan JSON. The plan schema is
    # arrays of fragments (key_facts/evidence_refs/data_gaps); pasting that in verbatim
    # demonstrates a bullet-list structure the model then mirrors into the final document.
    plan_json_text = _plan_to_brief(plan_json)

    prompt = f"""
You are a senior aerospace test/documentation engineer.
Generate a complete, realistic technical document in Markdown, based on the provided data and plan.

Date: {today}
Template type: {template_type}
Template description: {template_spec["description"]}
Template outline: {outline}
Writing goals: {template_spec["writing_goals"]}
Evidence style: {template_spec["evidence_style"]}
ECSS checklist (machine-readable):
{checklist_json}

Planning output:
{plan_json_text}

User description:
{description}

Additional notes:
{notes or "(none)"}

Uploaded file names:
{file_names or []}

Extracted evidence from uploaded files:
{uploaded_context or "(none)"}

{_drafting_rules(template_type)}
"""

    if progress_callback:
        progress_callback("Generating final document sections")

    if stream_draft_callback:
        if progress_callback:
            progress_callback("Starting live section streaming")
        # This overrides the JSON output contract above. Every other rule — drafting
        # form, tables, grounding — still applies and is the whole point of streaming
        # through the same prompt.
        stream_prompt = f"""
{prompt}

OVERRIDE — OUTPUT FORMAT FOR THIS RESPONSE ONLY:
Ignore the JSON output contract above. Output Markdown only, in this exact structure:
- First line: # <title>
- Then one ## heading per section, using this exact outline order and these exact heading strings: {outline_json}
- Do not add clause numbers to the ## headings themselves; number sub-headings inside a section as "### 4.1 ..." if useful.
- Fill each section with final polished content.

All other rules above still apply in full — especially the drafting form (prose vs lists),
the required Markdown tables, and the grounding/citation rules.
"""
        stream_completion = client.chat.completions.create(
            model=model_name,
            temperature=1,
            stream=True,
            messages=[
                {"role": "system", "content": "You generate precise engineering documents."},
                {"role": "user", "content": stream_prompt},
            ],
        )
        raw_markdown_parts: list[str] = []
        for chunk in stream_completion:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            piece = (delta.content or "") if delta else ""
            if not piece:
                continue
            raw_markdown_parts.append(piece)
            stream_draft_callback(piece)
        markdown_doc = "".join(raw_markdown_parts).strip()
        if progress_callback:
            progress_callback("Compiling streamed sections into template")
        parsed_markdown = _parse_markdown_sections(
            markdown_doc,
            outline,
            f"{template_type.replace('_', ' ').title()} Document",
        )
        validated = GeneratedDocument.model_validate(parsed_markdown)
        if progress_callback:
            progress_callback("Running final validation checks")
        return _validate_generated_document(validated.model_dump(), outline, template_type, plan_json)

    completion = client.chat.completions.create(
        model=model_name,
        temperature=1,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": "You generate precise engineering documents."},
            {"role": "user", "content": prompt},
        ],
    )
    raw = completion.choices[0].message.content or "{}"
    parsed = json.loads(raw)
    validated = GeneratedDocument.model_validate(parsed)
    return _validate_generated_document(validated.model_dump(), outline, template_type, plan_json)


# Templates whose content is legitimately step- or list-shaped. Everything else must
# default to continuous prose, which is what ECSS drafting conventions expect.
_LIST_SHAPED_TEMPLATES = {"test_procedure", "test_log"}


def _drafting_rules(template_type: str) -> str:
    """Output-form rules for the writer. Without these the model defaults to bullets."""
    if template_type in _LIST_SHAPED_TEMPLATES:
        form = """- Procedure/log steps are numbered lists, one action per step, imperative mood ("Apply 28 V to the primary bus.").
- Every list of steps is introduced by a full lead-in sentence ending in a colon, and followed by at least one sentence stating the expected outcome.
- Non-step sections (purpose, scope, prerequisites, safety, criteria) are written as continuous prose paragraphs, not bullets."""
    else:
        form = """- Write CONTINUOUS PROSE. Every section opens with at least two complete paragraphs of roughly 60-120 words each, in full sentences. This is the default and dominant form.
- A bulleted or numbered list is permitted ONLY for: enumerated applicable/reference documents, acronym definitions, or an explicitly ordered sequence. Introduce every list with a full lead-in sentence ending in a colon, and follow it with at least one sentence of interpretation.
- Never open a section with a list. Never let a section consist only of lists.
- NEVER emit bare sentence fragments as standalone lines."""

    return f"""Rules - OUTPUT CONTRACT:
- Return valid JSON only with this schema:
  {{
    "title": "string",
    "sections": [{{"title": "string", "content": "string"}}]
  }}
- Section titles must match the outline strings EXACTLY. Do not add clause numbers to them. Put any clause numbering INSIDE content as "### 4.1 ..." sub-headings.
- Do not include any keys other than title and sections.

Rules - DRAFTING FORM:
{form}
- Use third person, present tense, formal aerospace register. No second person. No contractions. No marketing adjectives.
- Normative statements use "shall"; recommendations "should"; permissions "may"; statements of fact "is/are".

Rules - TABLES (use GitHub-flavoured Markdown pipe tables):
- A verification/traceability/cross-reference section MUST contain a table with columns: | Req ID | Requirement | Verification Method (A/T/I/R) | Evidence Ref | Status |
- A results section MUST contain a table with columns: | Test ID | Parameter | Unit | Required | Measured | Verdict |
- An anomaly/non-conformance section MUST contain a table with columns: | NCR ID | Description | Severity | Affected Req | Disposition | Status |
- An interface section MUST contain a table with columns: | Interface ID | Type | Parameter | Value | Tolerance | Verification |
- Populate tables ONLY from supplied evidence. Where a value is genuinely absent write "TBD" (expected, not yet supplied) or "TBC" (supplied, not yet confirmed). Do NOT fabricate rows.

Rules - GROUNDING:
- Do not invent values, dates, part numbers, or measurements. You may assign structural identifiers (e.g. [REQ-EPS-010], section numbers) for organisation; the technical content they carry must come from the evidence.
- After each fact drawn from an uploaded file, cite the source inline as [filename].
- Where a required element has no supporting evidence, write a full sentence declaring the gap and its consequence, e.g. "No calibration record was supplied for the reference sensor; the accuracy claim in [REQ-TST-040] is therefore Not Verified." Do NOT write a bare "Not provided" line.
- Do not copy the brief verbatim; convert it into polished final prose.
- Ensure required_sections from the checklist are fully present in order."""


def _plan_to_brief(plan_json: dict) -> str:
    """
    Flatten the plan into a prose brief. Passing the raw plan JSON (arrays of short
    fragments) shows the writer a bullet-shaped template, which it then copies.
    """
    parts: list[str] = []
    for section in plan_json.get("sections", []):
        title = str(section.get("title", "")).strip()

        def _join(key: str) -> str:
            values = section.get(key, [])
            if not isinstance(values, list):
                return ""
            return "; ".join(str(v).strip() for v in values if str(v).strip())

        facts = _join("key_facts")
        refs = _join("evidence_refs")
        gaps = _join("data_gaps")
        acceptance = str(section.get("acceptance_statement", "")).strip()

        parts.append(
            f"SECTION: {title}\n"
            f"Material to work into flowing paragraphs: {facts or 'none'}\n"
            f"Sources to cite inline: {refs or 'none'}\n"
            f"Declare these as unavailable, in full sentences: {gaps or 'none'}\n"
            f"Acceptance intent: {acceptance or 'none'}"
        )
    return "\n\n".join(parts) if parts else "(no plan available)"


def _normalize_heading(s: str) -> str:
    """
    Fuzzy key for matching an outline title against whatever heading the model wrote.
    Models routinely Title-Case, add clause numbers, or expand "&" -> "and"; none of
    those should cost us the section body.
    """
    s = str(s).strip().lower()
    s = re.sub(r"^\d+(\.\d+)*[.)]?\s*", "", s)  # leading "4." / "2.3)" clause numbers
    s = s.replace("&", " and ")
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return " ".join(s.split())


def _validate_generated_document(doc: dict, outline: list[str], template_type: str, plan_json: dict) -> dict:
    """
    Lightweight post-generation validation:
    - enforce exact section coverage/order from outline
    - prevent empty sections
    - reduce obvious "Not provided" misuse
    """
    raw_sections = [
        (str(s.get("title", "")).strip(), str(s.get("content", "")).strip())
        for s in doc.get("sections", [])
        if str(s.get("title", "")).strip()
    ]

    planned_map: dict[str, str] = {}
    for section in plan_json.get("sections", []):
        title = str(section.get("title", "")).strip()
        gaps = section.get("data_gaps", [])
        if title and isinstance(gaps, list) and gaps:
            joined = "; ".join(str(g).strip() for g in gaps if str(g).strip())
            if joined:
                planned_map[_normalize_heading(title)] = (
                    f"The following required information was not supplied in the provided inputs: {joined}."
                )

    # Index model sections by normalized heading. Keep them as ordered queues so a
    # duplicate heading from the model resolves to distinct sections, and so a section
    # consumed by a key match can never be reused by the positional fallback.
    by_key: dict[str, list[int]] = {}
    for idx, (title, content) in enumerate(raw_sections):
        if content:
            by_key.setdefault(_normalize_heading(title), []).append(idx)

    resolved: list[Optional[str]] = [None] * len(outline)
    used: set[int] = set()

    # Pass 1 — match each outline title to a model heading (case / clause-number / &-and).
    for pos, required_title in enumerate(outline):
        queue = by_key.get(_normalize_heading(required_title))
        if queue:
            idx = queue.pop(0)
            used.add(idx)
            resolved[pos] = raw_sections[idx][1]

    # Pass 2 — the model kept content but renamed the heading past what normalization
    # tolerates. Fill still-empty outline slots from the leftover sections in order,
    # but only when it's an unambiguous 1:1 so we never scramble sections.
    empty_positions = [pos for pos in range(len(outline)) if resolved[pos] is None]
    leftover = [idx for idx in range(len(raw_sections)) if idx not in used and raw_sections[idx][1]]
    if empty_positions and len(empty_positions) == len(leftover):
        for pos, idx in zip(empty_positions, leftover):
            resolved[pos] = raw_sections[idx][1]
            logger.warning(
                "Section %r filled positionally from model section %r",
                outline[pos],
                raw_sections[idx][0],
            )

    normalized_sections = []
    for pos, required_title in enumerate(outline):
        content = resolved[pos]
        if not content:
            content = planned_map.get(_normalize_heading(required_title), "Not provided")
            logger.warning("Section %r had no model content; fell back to plan/placeholder", required_title)
        normalized_sections.append({"title": required_title, "content": content})

    title = str(doc.get("title", "")).strip() or f"{template_type.replace('_', ' ').title()} Document"
    return {"title": title, "sections": normalized_sections}


def _sse_event(event: str, data: dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


@app.post("/generate-document")
async def generate_document(
    description: str = Form(...),
    template_type: str = Form(...),
    notes: str = Form(""),
    section_outline: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
):
    parsed_outline: Optional[List[str]] = None
    if section_outline:
        try:
            raw_outline = json.loads(section_outline)
            if isinstance(raw_outline, list):
                parsed_outline = [str(x) for x in raw_outline]
        except Exception:
            parsed_outline = None

    try:
        client, model_name = _get_llm_client()
        file_names, uploaded_context = await _build_uploaded_context(files, client, model_name)
        return _generate_doc_from_inputs(
            client=client,
            model_name=model_name,
            description=description,
            notes=notes,
            template_type=template_type,
            section_outline=parsed_outline,
            file_names=file_names,
            uploaded_context=uploaded_context,
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"AI generation failed: {exc}") from exc


@app.post("/generate-document-stream")
async def generate_document_stream(
    description: str = Form(...),
    template_type: str = Form(...),
    notes: str = Form(""),
    section_outline: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
):
    parsed_outline: Optional[List[str]] = None
    if section_outline:
        try:
            raw_outline = json.loads(section_outline)
            if isinstance(raw_outline, list):
                parsed_outline = [str(x) for x in raw_outline]
        except Exception:
            parsed_outline = None

    async def event_stream() -> AsyncGenerator[str, None]:
        try:
            yield _sse_event("progress", {"message": "Preparing generation context"})
            client, model_name = _get_llm_client()
            yield _sse_event("progress", {"message": "Connection ready. Starting evidence extraction"})
            loop = asyncio.get_running_loop()
            draft_queue: asyncio.Queue[str | None] = asyncio.Queue()
            progress_queue: asyncio.Queue[str | None] = asyncio.Queue()
            partial_chunks: list[str] = []
            first_chunk_ready = asyncio.Event()
            early_plan_task: Optional[asyncio.Task] = None

            def _on_draft_chunk(chunk_text: str) -> None:
                loop.call_soon_threadsafe(draft_queue.put_nowait, chunk_text)

            def _on_progress(message: str) -> None:
                loop.call_soon_threadsafe(progress_queue.put_nowait, message)

            def _on_chunk_ready(file_name: str, text: str) -> None:
                partial_chunks.append(text)
                if not first_chunk_ready.is_set():
                    first_chunk_ready.set()
                loop.call_soon_threadsafe(progress_queue.put_nowait, f"Evidence extracted: {file_name}")

            outline = _resolve_outline(template_type, parsed_outline)
            template_spec = TEMPLATE_PROMPT_SPECS.get(template_type, TEMPLATE_PROMPT_SPECS["test_procedure"])
            today = date.today().strftime("%d/%m/%Y")

            extraction_task = asyncio.create_task(
                _build_uploaded_context(
                    files,
                    client,
                    model_name,
                    on_chunk_ready=_on_chunk_ready,
                )
            )

            while not extraction_task.done():
                while not progress_queue.empty():
                    progress = progress_queue.get_nowait()
                    if progress:
                        yield _sse_event("progress", {"message": progress})

                if first_chunk_ready.is_set() and early_plan_task is None:
                    early_context = "\n\n".join(partial_chunks).strip()
                    if early_context:
                        yield _sse_event("progress", {"message": "Starting early planning from first evidence chunk"})
                        early_plan_task = asyncio.create_task(
                            asyncio.to_thread(
                                _plan_from_context,
                                client,
                                model_name,
                                template_type,
                                outline,
                                template_spec,
                                today,
                                description,
                                notes,
                                [f.filename or "upload" for f in files],
                                early_context,
                                None,
                            )
                        )
                await asyncio.sleep(0.04)

            file_names, uploaded_context = await extraction_task
            yield _sse_event("progress", {"message": "All evidence processed. Building final context"})

            early_plan_json: Optional[dict[str, Any]] = None
            if early_plan_task:
                try:
                    early_plan_json = await early_plan_task
                    yield _sse_event("progress", {"message": "Early plan complete. Refining with full context"})
                except Exception:
                    early_plan_json = None
                    yield _sse_event("progress", {"message": "Early plan unavailable. Continuing with full-context plan"})

            final_future = asyncio.to_thread(
                _generate_doc_from_inputs,
                client,
                model_name,
                description,
                notes,
                template_type,
                parsed_outline,
                file_names,
                uploaded_context,
                _on_progress,
                _on_draft_chunk,
                early_plan_json,
            )

            task = asyncio.create_task(final_future)

            while True:
                while not progress_queue.empty():
                    progress = progress_queue.get_nowait()
                    if progress:
                        yield _sse_event("progress", {"message": progress})

                while not draft_queue.empty():
                    chunk = draft_queue.get_nowait()
                    if chunk:
                        yield _sse_event("draft", {"text": chunk})

                if task.done():
                    break
                await asyncio.sleep(0.04)

            final_doc = await task
            while not draft_queue.empty():
                chunk = draft_queue.get_nowait()
                if chunk:
                    yield _sse_event("draft", {"text": chunk})
            yield _sse_event("completed", {"document": final_doc})
        except HTTPException as exc:
            yield _sse_event("error", {"message": str(exc.detail)})
        except Exception as exc:
            yield _sse_event("error", {"message": f"AI generation failed: {exc}"})

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@app.post("/compatibility-check")
async def compatibility_check(
    component_a: UploadFile = File(...),
    component_b: UploadFile = File(...),
):
    try:
        client, model_name = _get_llm_client()
        raw_a = await component_a.read()
        raw_b = await component_b.read()
        if not raw_a or not raw_b:
            raise HTTPException(status_code=400, detail="Both component files are required.")

        text_a = _extract_text_from_upload(component_a.filename or "component_a", component_a.content_type, raw_a).strip()
        text_b = _extract_text_from_upload(component_b.filename or "component_b", component_b.content_type, raw_b).strip()

        # Vision fallback for image-only inputs.
        if not text_a and _is_probably_image(component_a.filename or "component_a", component_a.content_type):
            text_a = _summarize_image_evidence(client, model_name, component_a.filename or "component_a", component_a.content_type, raw_a)
        if not text_b and _is_probably_image(component_b.filename or "component_b", component_b.content_type):
            text_b = _summarize_image_evidence(client, model_name, component_b.filename or "component_b", component_b.content_type, raw_b)

        if not text_a or not text_b:
            raise HTTPException(
                status_code=400,
                detail="Could not extract enough readable content from one or both files.",
            )

        prompt = f"""
You are an aerospace compatibility analyst.
Compare two component technical documents and return JSON only with this schema:
{{
  "<Parameter_Name>": {{
    "A": "value from component A",
    "B": "value from component B",
    "Status": "OK" | "Different" | "Missing"
  }},
  "Summary": "short overall compatibility summary"
}}
Rules:
- Include practical parameters for integration checks: supply voltage, power, interface/protocol, connector/pinout, data rate, operating temp, mass/dimensions if available.
- Use "Missing" when one or both values are absent.
- No arrays, no extra keys.

Component A ({component_a.filename}):
---
{text_a[:24000]}
---

Component B ({component_b.filename}):
---
{text_b[:24000]}
---
"""
        completion = client.chat.completions.create(
            model=model_name,
            temperature=1,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You compare engineering component specifications."},
                {"role": "user", "content": prompt},
            ],
        )
        content = completion.choices[0].message.content or "{}"
        parsed = json.loads(content)
        try:
            validated = validate_llm_payload(parsed)
        except ValueError as exc:
            raise HTTPException(status_code=502, detail=f"LLM response failed validation: {exc}") from exc
        return validated
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Compatibility check failed: {exc}") from exc


@app.post("/llm-compare-pdfs")
async def llm_compare_pdfs(
    pdf_a: UploadFile = File(...),
    pdf_b: UploadFile = File(...),
):
    def _maybe_debug_write(name: str, payload: str | bytes) -> None:
        out = os.getenv("DATA_DIR", "").strip()
        if not out:
            return
        try:
            os.makedirs(out, exist_ok=True)
            path = os.path.join(out, name)
            if isinstance(payload, bytes):
                with open(path, "wb") as fh:
                    fh.write(payload)
            else:
                with open(path, "w", encoding="utf-8") as fh:
                    fh.write(payload)
        except OSError:
            pass

    if not (pdf_a.filename or "").lower().endswith(".pdf") or not (pdf_b.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    try:
        raw_a = await pdf_a.read()
        raw_b = await pdf_b.read()
        text_a = extract_text_from_pdf(raw_a)
        text_b = extract_text_from_pdf(raw_b)
        _maybe_debug_write("debug_pdf_a.txt", text_a or "No text extracted")
        _maybe_debug_write("debug_pdf_b.txt", text_b or "No text extracted")
        if not text_a.strip() or not text_b.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from one or both PDFs")

        client, model_name = _get_llm_client()
        detected_a = _detect_component_type(client, model_name, pdf_a.filename or "A", text_a)
        detected_b = _detect_component_type(client, model_name, pdf_b.filename or "B", text_b)
        pair_key = _pair_key(detected_a["component_type"], detected_b["component_type"])
        pair_specific = PAIR_PARAMETER_WHITELISTS.get(pair_key)
        pair_supported = pair_specific is not None
        parameter_whitelist = pair_specific or UNIVERSAL_PARAMETER_WHITELIST
        parameter_hint = ", ".join(parameter_whitelist)
        completion = client.chat.completions.create(
            model=model_name,
            temperature=1,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": LLM_COMPARE_SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": f"""
Detected component A type: {detected_a["component_type"]} (confidence {detected_a["confidence"]:.2f})
Detected component B type: {detected_b["component_type"]} (confidence {detected_b["confidence"]:.2f})
Pair key: {pair_key}
Pair supported by focused rules: {"yes" if pair_supported else "no (fallback to universal core fields)"}

Important:
- Compare ONLY these parameter names if present in docs: {parameter_hint}
- Do NOT include generic identity metadata like product_name, website, contact_email, version unless strictly needed for compatibility.
- Keep output compact and compatibility-focused.
- For interface-like rows (interface/protocol/connector), if both components share at least one common interface, mark Status as "OK" even if one side supports additional interfaces.
- For input voltage differences, keep the difference but mention in Summary when compatibility may still be feasible via external DC/DC conditioning.
- Avoid direct compatibility verdicts for mass/dimensions unless the selected pair explicitly needs structural fit checks.

Product A:
{text_a[:48000]}

Product B:
{text_b[:48000]}
""",
                },
            ],
        )
        content = completion.choices[0].message.content or "{}"
        raw = json.loads(content)
        if isinstance(raw, dict):
            raw["DetectedComponents"] = {
                "component_a_type": detected_a["component_type"],
                "component_b_type": detected_b["component_type"],
                "component_a_confidence": detected_a["confidence"],
                "component_b_confidence": detected_b["confidence"],
                "pair_key": pair_key,
                "pair_supported": pair_supported,
                "component_a_evidence": detected_a["evidence"],
                "component_b_evidence": detected_b["evidence"],
            }
            raw = _apply_compatibility_overrides(raw, pair_key)
        try:
            result = validate_llm_payload(raw)
        except ValueError as exc:
            _maybe_debug_write("llm_output_invalid.json", json.dumps(raw, indent=2))
            raise HTTPException(status_code=502, detail=f"LLM response failed schema validation: {exc}") from exc
        _maybe_debug_write("llm_output.txt", json.dumps(result, indent=2))
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF comparison failed: {exc}") from exc


class AnalyzedTemplate(BaseModel):
    outline: List[str]
    notes: str = ""


def _normalize_text_for_pdf(value: str) -> str:
    """
    Preserve original symbols; only strip control chars that can break rendering.
    """
    # Remove control chars that can break rendering, keep newlines/tabs.
    out = re.sub(r"[^\x09\x0A\x0D\x20-\uFFFF]", "", value)
    return out


def _safe_ascii_filename(name: str) -> str:
    clean = _normalize_text_for_pdf(name)
    clean = clean.encode("ascii", "ignore").decode("ascii")
    clean = re.sub(r'[^A-Za-z0-9._ -]+', "", clean).strip()
    return clean or "document"


def _register_pdf_font() -> str:
    """
    Prefer a Unicode-capable TTF font; fallback to Helvetica if unavailable.
    """
    candidates = [
        ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ("Arial", "C:/Windows/Fonts/arial.ttf"),
        ("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"),
    ]
    for font_name, font_path in candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                return font_name
            except Exception:
                continue
    return "Helvetica"


@app.post("/export-pdf")
def export_pdf(doc: GeneratedDocument):
    styles = getSampleStyleSheet()
    font_name = _register_pdf_font()
    styles["Title"].fontName = font_name
    styles["Heading2"].fontName = font_name
    styles["BodyText"].fontName = font_name
    title = _normalize_text_for_pdf(doc.title)
    safe_filename = _safe_ascii_filename(doc.title)

    margin = 36
    frame_width = A4[0] - 2 * margin - 8  # small slack for grid lines

    def _render(tables_as_text: bool) -> bytes:
        buffer = io.BytesIO()
        pdf = SimpleDocTemplate(
            buffer, pagesize=A4, leftMargin=margin, rightMargin=margin,
            topMargin=margin, bottomMargin=margin, title=title,
        )
        elements = [Paragraph(title, styles["Title"]), Spacer(1, 14)]
        for section in doc.sections:
            elements.append(Paragraph(_normalize_text_for_pdf(section.title), styles["Heading2"]))
            elements.append(Spacer(1, 6))
            clean = _normalize_text_for_pdf(section.content or "")
            elements.extend(
                _markdown_block_to_pdf_flowables(clean, styles, font_name, frame_width, tables_as_text)
            )
            elements.append(Spacer(1, 12))
        pdf.build(elements)
        return buffer.getvalue()

    try:
        data = _render(tables_as_text=False)
    except Exception as exc:
        # A single table row taller than the page can't be split and aborts the
        # build. Retry with tables flattened to text so export never 500s on shape.
        logger.warning("PDF table layout failed (%s); retrying with tables as text", exc)
        try:
            data = _render(tables_as_text=True)
        except Exception as exc2:
            raise HTTPException(status_code=500, detail=f"PDF export failed: {exc2}") from exc2

    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}.pdf"'},
    )


# Inline markdown we care about: **bold**, *italic*/_italic_, `code`.
_MD_INLINE_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*\n]+?\*|_[^_\n]+?_|`[^`\n]+?`)")

_MD_TABLE_SEP_RE = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def _md_inline_to_rl(text: str) -> str:
    """Convert inline markdown to ReportLab's mini-HTML, escaping everything else."""
    out: list[str] = []
    for part in _MD_INLINE_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            out.append(f"<b>{escape(part[2:-2])}</b>")
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            out.append(f"<i>{escape(part[1:-1])}</i>")
        elif part.startswith("`") and part.endswith("`"):
            out.append(f"<font face='Courier'>{escape(part[1:-1])}</font>")
        else:
            out.append(escape(part))
    return "".join(out)


def _markdown_block_to_pdf_flowables(
    content: str, styles, font_name: str, frame_width: float = 515.0, tables_as_text: bool = False
) -> list:
    """
    Render a section body to ReportLab flowables. The generation prompt mandates GFM
    pipe tables, so they must become real tables here rather than rows of "|".

    tables_as_text renders tables as plain lines instead of Table flowables — the
    fallback path when a genuine table would overflow a page and abort the build.
    """
    flowables: list = []
    lines = (content or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # GFM pipe table: header row, then a |---|---| separator.
        if line.startswith("|") and i + 1 < len(lines) and _MD_TABLE_SEP_RE.match(lines[i + 1]):
            header = [c.strip() for c in line.strip("|").split("|")]
            rows = [header]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            width = max(1, len(header))

            if tables_as_text:
                for r_idx, row in enumerate(rows):
                    padded = (row + [""] * (width - len(row)))[:width]
                    flowables.append(Paragraph(_md_inline_to_rl(" | ".join(padded)), styles["BodyText"]))
                    if r_idx == 0:
                        flowables.append(Spacer(1, 2))
                flowables.append(Spacer(1, 8))
                continue

            cell_style = ParagraphStyle(
                "TableCell", parent=styles["BodyText"], fontName=font_name, fontSize=7.5, leading=9.5
            )
            data = [
                [Paragraph(_md_inline_to_rl(c), cell_style) for c in (row + [""] * (width - len(row)))[:width]]
                for row in rows
            ]
            # Fixed, equal column widths that sum to the frame so wrapping is
            # predictable and no column is starved.
            col_widths = [frame_width / width] * width
            table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            flowables.append(table)
            flowables.append(Spacer(1, 8))
            continue

        heading = re.match(r"^(#{3,6})\s+(.*)$", line)
        if heading:
            style = styles["Heading3"]
            style.fontName = font_name
            flowables.append(Paragraph(_md_inline_to_rl(heading.group(2).strip()), style))
            flowables.append(Spacer(1, 4))
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", line)
        numbered = re.match(r"^(\d+[.)])\s+(.*)$", line)
        if bullet or numbered:
            marker = "•" if bullet else numbered.group(1)
            body = bullet.group(1) if bullet else numbered.group(2)
            item_style = ParagraphStyle(
                "MdListItem", parent=styles["BodyText"], fontName=font_name, leftIndent=14, bulletIndent=4
            )
            flowables.append(Paragraph(_md_inline_to_rl(body), item_style, bulletText=marker))
            i += 1
            continue

        # Gather a paragraph until a blank line or a block-level marker.
        para: list[str] = []
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("|") or re.match(r"^(#{3,6}\s|[-*+]\s|\d+[.)]\s)", nxt):
                break
            para.append(nxt)
            i += 1
        if para:
            flowables.append(Paragraph(_md_inline_to_rl(" ".join(para)), styles["BodyText"]))
            flowables.append(Spacer(1, 6))
    return flowables


def _add_markdown_runs(paragraph, text: str) -> None:
    """Split a line on inline markdown markers and add styled runs to a docx paragraph."""
    for part in _MD_INLINE_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def _render_markdown_block_to_docx(document, content: str) -> None:
    """Map a section's markdown-ish body onto Word paragraphs, lists and tables."""
    lines = (content or "").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # Markdown table: a header row followed by a |---|---| separator.
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip("|").split("|")]
            rows: List[List[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for idx, cell_text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(min(len(row), len(header))):
                    cells[idx].text = ""
                    _add_markdown_runs(cells[idx].paragraphs[0], row[idx])
            document.add_paragraph()
            continue

        # Nested sub-heading inside a section body.
        heading = re.match(r"^(#{3,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 6)
            document.add_heading(heading.group(2).strip(), level=level)
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", line)
        if bullet:
            para = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(para, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.*)$", line)
        if numbered:
            para = document.add_paragraph(style="List Number")
            _add_markdown_runs(para, numbered.group(1))
            i += 1
            continue

        para = document.add_paragraph()
        _add_markdown_runs(para, line)
        i += 1


@app.post("/export-docx")
def export_docx(doc: GeneratedDocument):
    try:
        document = Document()
        document.add_heading(doc.title or "Document", level=0)

        for section in doc.sections:
            document.add_heading(section.title, level=1)
            _render_markdown_block_to_docx(document, section.content or "")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"{_safe_ascii_filename(doc.title)}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {exc}") from exc


def _extract_pdf_text(data: bytes, max_pages: int = 20, max_chars: int = 24_000) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: List[str] = []
    for i, page in enumerate(reader.pages[:max_pages]):
        t = page.extract_text() or ""
        parts.append(t)
    text = "\n\n".join(parts).strip()
    return text[:max_chars]


def _is_probably_image(filename: str, content_type: str | None) -> bool:
    ct = (content_type or "").lower()
    if ct.startswith("image/"):
        return True
    low = filename.lower()
    return low.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif"))


@app.post("/analyze-template")
async def analyze_template(
    file: UploadFile = File(...),
    document_category: str = Form(...),
):
    """
    Infer section headings from a company template (PDF with extractable text, or image/screenshot).
    The user selects `document_category` (one of the five built-in template_type ids) for context.
    """
    if document_category not in TEMPLATE_OUTLINES:
        raise HTTPException(
            status_code=400,
            detail=f"document_category must be one of: {', '.join(TEMPLATE_OUTLINES.keys())}",
        )

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file")

    reference_outline = TEMPLATE_OUTLINES[document_category]
    filename = file.filename or "upload"

    user_content: list[dict]

    if _is_probably_image(filename, file.content_type):
        mime = file.content_type or "image/png"
        if not mime.startswith("image/"):
            mime = "image/png"
        b64 = base64.standard_b64encode(raw).decode("ascii")
        user_content = [
            {
                "type": "text",
                "text": f"""This is an image of a company document template or a filled example.
Document category (for context only): {document_category}
Typical sections for this category are: {reference_outline}

List the main section or chapter titles you see in the template (table of contents, headings, numbered clauses).
Return JSON only: {{"outline": ["string", ...], "notes": "optional short note"}}
Rules:
- 4–20 sections, short titles as they appear (or normalized Title Case).
- Preserve order top-to-bottom.
- Do not include page numbers-only lines.""",
            },
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
        ]
    else:
        text = _extract_pdf_text(raw)
        if len(text) < 80:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not read enough text from this PDF (it may be scanned). "
                    "Export the template as a PNG/JPG screenshot or use a text-based PDF."
                ),
            )
        user_content = [
            {
                "type": "text",
                "text": f"""Below is extracted text from a company document template or example (PDF).
Document category (for context): {document_category}
Typical sections for this category: {reference_outline}

Extract the main section / chapter headings. Return JSON only:
{{"outline": ["string", ...], "notes": "optional"}}

Text:
---
{text}
---""",
            }
        ]

    try:
        client, model_name = _get_llm_client()
        completion = client.chat.completions.create(
            model=model_name,
            temperature=1,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "You read engineering documents and extract structured headings. Return JSON only.",
                },
                {"role": "user", "content": user_content},
            ],
        )
        raw_json = completion.choices[0].message.content or "{}"
        parsed = json.loads(raw_json)
        validated = AnalyzedTemplate.model_validate(parsed)
        if not validated.outline:
            raise ValueError("empty outline")
        return validated.model_dump()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Template analysis failed: {exc}") from exc
